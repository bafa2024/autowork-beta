#!/usr/bin/env python3
"""
Project Management API Server
Provides REST API endpoints for the project management dashboard
"""

import os
import json
import logging
import io
from datetime import datetime, timedelta
from typing import Dict, List, Optional
from flask import Flask, request, jsonify, render_template_string, send_file
from flask_cors import CORS
from dotenv import load_dotenv

# Import project management system
from project_management import ProjectManager, ProjectStatus, TaskStatus, Priority

# Load environment variables
load_dotenv()

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__)
CORS(app)  # Enable CORS for frontend

# Initialize project manager
freelancer_token = os.environ.get('FREELANCER_OAUTH_TOKEN')
if not freelancer_token:
    logger.error("FREELANCER_OAUTH_TOKEN not found in environment variables")
    exit(1)

project_manager = ProjectManager(
    freelancer_token=freelancer_token,
    redis_url=os.environ.get('REDIS_URL')
)

@app.route('/')
def index():
    """Serve the project dashboard"""
    try:
        with open('project_dashboard.html', 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return "Project dashboard HTML file not found", 404

@app.route('/api/dashboard', methods=['GET'])
def get_dashboard():
    """Get dashboard data"""
    try:
        data = project_manager.get_dashboard_data()
        return jsonify(data)
    except Exception as e:
        logger.error(f"Error getting dashboard data: {e}")
        return jsonify({
            'error': str(e),
            'metrics': {
                'active_projects': 0,
                'total_hours_tracked': 0,
                'completion_rate': 0
            },
            'active_projects': [],
            'today_tasks': [],
            'upcoming_deadlines': []
        }), 500

@app.route('/api/projects', methods=['GET'])
def get_projects():
    """Get all projects"""
    try:
        from project_management import ManagedProject
        projects = project_manager.session.query(ManagedProject).all()
        return jsonify([project_manager._project_summary(p) for p in projects])
    except Exception as e:
        logger.error(f"Error getting projects: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/projects/sync', methods=['POST'])
def sync_projects():
    """Sync projects from Freelancer"""
    try:
        imported_projects = project_manager.fetch_and_import_awarded_projects()
        return jsonify({
            'success': True,
            'imported_count': len(imported_projects),
            'message': f'Imported {len(imported_projects)} new projects'
        })
    except Exception as e:
        logger.error(f"Error syncing projects: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/projects/import', methods=['POST'])
def import_project():
    """Import a specific project by ID"""
    try:
        data = request.get_json()
        project_id = data.get('freelancer_project_id')
        
        if not project_id:
            return jsonify({
                'success': False,
                'error': 'Project ID is required'
            }), 400
        
        # Check if project already exists
        from project_management import ManagedProject
        existing = project_manager.session.query(ManagedProject).filter(
            ManagedProject.freelancer_project_id == project_id
        ).first()
        
        if existing:
            return jsonify({
                'success': False,
                'error': 'Project already exists in system'
            }), 400
        
        # Import the project
        project = project_manager.import_awarded_project(project_id)
        
        if project:
            return jsonify({
                'success': True,
                'project': project_manager._project_summary(project),
                'message': f'Project {project_id} imported successfully'
            })
        else:
            return jsonify({
                'success': False,
                'error': f'Failed to import project {project_id}'
            }), 400
            
    except Exception as e:
        logger.error(f"Error importing project: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/projects/<int:project_id>', methods=['GET'])
def get_project(project_id):
    """Get specific project details"""
    try:
        from project_management import ManagedProject
        project = project_manager.session.query(ManagedProject).filter(
            ManagedProject.id == project_id
        ).first()
        
        if not project:
            return jsonify({'error': 'Project not found'}), 404
        
        # Get detailed project data
        project_data = project_manager._project_summary(project)
        
        # Add tasks
        project_data['tasks'] = [project_manager._task_summary(task) for task in project.tasks]
        
        # Add recent communications
        recent_comms = project_manager.session.query(project_manager.Communication).filter(
            project_manager.Communication.project_id == project_id
        ).order_by(project_manager.Communication.timestamp.desc()).limit(5).all()
        
        project_data['recent_communications'] = [{
            'subject': comm.subject,
            'type': comm.message_type,
            'timestamp': comm.timestamp.isoformat(),
            'content': comm.content[:100] + '...' if len(comm.content) > 100 else comm.content
        } for comm in recent_comms]
        
        # Add risk assessment
        risk_assessment = project_manager.assess_project_risk(project_id)
        project_data['risk_assessment'] = risk_assessment
        
        return jsonify(project_data)
        
    except Exception as e:
        logger.error(f"Error getting project {project_id}: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/projects/<int:project_id>/status', methods=['PUT'])
def update_project_status(project_id):
    """Update project status"""
    try:
        data = request.get_json()
        new_status = data.get('status')
        
        if not new_status:
            return jsonify({'error': 'Status is required'}), 400
        
        # Validate status
        try:
            status_enum = ProjectStatus(new_status)
        except ValueError:
            return jsonify({'error': f'Invalid status: {new_status}'}), 400
        
        success = project_manager.update_project_status(project_id, status_enum)
        
        if success:
            return jsonify({'success': True, 'message': 'Project status updated'})
        else:
            return jsonify({'error': 'Failed to update project status'}), 400
            
    except Exception as e:
        logger.error(f"Error updating project status: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/tasks/<int:task_id>/status', methods=['PUT'])
def update_task_status(task_id):
    """Update task status"""
    try:
        data = request.get_json()
        new_status = data.get('status')
        
        if not new_status:
            return jsonify({'error': 'Status is required'}), 400
        
        # Validate status
        try:
            status_enum = TaskStatus(new_status)
        except ValueError:
            return jsonify({'error': f'Invalid status: {new_status}'}), 400
        
        success = project_manager.update_task_status(task_id, status_enum)
        
        if success:
            # Recalculate project progress
            task = project_manager.session.query(project_manager.Task).filter(
                project_manager.Task.id == task_id
            ).first()
            if task:
                project_manager.recalculate_project_progress(task.project_id)
            
            return jsonify({'success': True, 'message': 'Task status updated'})
        else:
            return jsonify({'error': 'Failed to update task status'}), 400
            
    except Exception as e:
        logger.error(f"Error updating task status: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/projects/<int:project_id>/time', methods=['POST'])
def log_time(project_id):
    """Log time for a project"""
    try:
        data = request.get_json()
        hours = data.get('hours')
        description = data.get('description')
        
        if not hours or not description:
            return jsonify({'error': 'Hours and description are required'}), 400
        
        if hours <= 0:
            return jsonify({'error': 'Hours must be positive'}), 400
        
        project_manager.log_time(
            project_id=project_id,
            task_id=None,  # Could be made optional
            hours=hours,
            description=description,
            user="Developer"
        )
        
        # Recalculate project progress
        project_manager.recalculate_project_progress(project_id)
        
        return jsonify({'success': True, 'message': 'Time logged successfully'})
        
    except Exception as e:
        logger.error(f"Error logging time: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/projects/<int:project_id>/report', methods=['GET'])
def get_project_report(project_id):
    """Generate project report"""
    try:
        from project_management import ManagedProject
        project = project_manager.session.query(ManagedProject).filter(
            ManagedProject.id == project_id
        ).first()
        
        if not project:
            return jsonify({'error': 'Project not found'}), 404
        
        # Generate report
        report_text = project_manager.generate_project_report(project_id)
        
        # Get project data for the report
        project_data = project_manager._project_summary(project)
        tasks = [project_manager._task_summary(task) for task in project.tasks]
        
        return jsonify({
            'success': True,
            'project_title': project.title,
            'status': project.status.value,
            'progress': project.progress_percentage,
            'hours_tracked': project.hours_tracked,
            'deadline': project.deadline.isoformat() if project.deadline else None,
            'tasks': tasks,
            'recent_activity': report_text
        })
        
    except Exception as e:
        logger.error(f"Error generating report: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/projects/<int:project_id>/client-update', methods=['POST'])
def send_client_update(project_id):
    """Send client update"""
    try:
        data = request.get_json()
        update_type = data.get('type', 'progress')
        
        project_manager.send_client_update(project_id, update_type)
        
        return jsonify({'success': True, 'message': 'Client update sent successfully'})
        
    except Exception as e:
        logger.error(f"Error sending client update: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    from project_management import ManagedProject
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'project_count': project_manager.session.query(ManagedProject).count()
    })

@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Endpoint not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/sdlc/analyze', methods=['POST'])
def analyze_project_sdlc():
    """Analyze project and generate complete SDLC documents including task breakdown and versioned releases"""
    try:
        data = request.get_json()
        project_description = data.get('project_description', '')
        project_title = data.get('project_title', 'Project')
        
        if not project_description:
            return jsonify({
                'success': False,
                'error': 'Project description is required'
            }), 400
        
        # Import SDLC service
        from auto_sdlc_service import AutoSDLCService
        
        sdlc_service = AutoSDLCService()
        
        # Generate complete SDLC documents including task breakdown and versioned releases
        result = sdlc_service.generate_complete_sdlc(project_description)
        
        # Prepare response with additional metadata
        response_data = {
            'success': True,
            'summary': {
                'project_type': result['project_analysis']['project_type'],
                'complexity': result['project_analysis']['complexity'],
                'estimated_hours': result['project_analysis']['estimated_hours'],
                'technologies': result['project_analysis']['technologies'],
                'key_features': result['project_analysis']['key_features'],
                'risks': result['project_analysis']['risks'],
                'total_tasks': result['metadata']['total_tasks'],
                'total_versions': result['metadata']['total_versions']
            },
            'srs': result['srs_document'],
            'design': result['system_design'],
            'ui_design': result['ui_design'],
            'implementation_plan': result['implementation_plan'],
            'implementation_tools': result['implementation_tools'],
            'task_breakdowns': result['task_breakdowns'],
            'versioned_releases': result['versioned_releases'],
            'test_plan': result['test_plan'],
            'deployment_plan': result['deployment_plan'],
            'maintenance_plan': result['maintenance_plan'],
            'coding_prompts': result.get('coding_prompts', []),
            'testing_prompts': result.get('testing_prompts', []),
            'metadata': result['metadata']
        }
        
        return jsonify(response_data)
        
    except Exception as e:
        logger.error(f"Error analyzing project SDLC: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/sdlc/export', methods=['POST'])
def export_sdlc_documents():
    """Export SDLC documents in various formats"""
    try:
        data = request.get_json()
        srs_data = data.get('srs', {})
        design_data = data.get('design', {})
        ui_design_data = data.get('ui_design', {})
        task_breakdowns_data = data.get('task_breakdowns', [])
        versioned_releases_data = data.get('versioned_releases', [])
        plan_data = data.get('implementation_plan', {})
        implementation_tools_data = data.get('implementation_tools', {})
        test_plan_data = data.get('test_plan', {})
        deployment_plan_data = data.get('deployment_plan', {})
        maintenance_plan_data = data.get('maintenance_plan', {})
        export_format = data.get('format', 'json')
        project_title = data.get('project_title', 'Project')
        
        # Import SDLC service
        from auto_sdlc_service import AutoSDLCService, SRSDocument, DesignDocument, UIDesignDocument, ImplementationPlan
        
        sdlc_service = AutoSDLCService()
        
        # Reconstruct objects from data
        srs = SRSDocument(**srs_data)
        design = DesignDocument(**design_data)
        ui_design = UIDesignDocument(**ui_design_data)
        plan = ImplementationPlan(**plan_data)
        
        if export_format in ['word', 'pdf']:
            # Generate document in memory
            if export_format == 'word':
                doc_buffer = generate_word_document(srs, design, ui_design, task_breakdowns_data, versioned_releases_data, plan, implementation_tools_data, test_plan_data, deployment_plan_data, maintenance_plan_data, project_title)
                filename = f"{project_title.replace(' ', '_')}_Complete_SDLC_Documents.docx"
                mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:  # pdf
                doc_buffer = generate_pdf_document(srs, design, ui_design, task_breakdowns_data, versioned_releases_data, plan, implementation_tools_data, test_plan_data, deployment_plan_data, maintenance_plan_data, project_title)
                filename = f"{project_title.replace(' ', '_')}_Complete_SDLC_Documents.pdf"
                mimetype = 'application/pdf'
            
            # Return file for download
            doc_buffer.seek(0)
            return send_file(
                doc_buffer,
                mimetype=mimetype,
                as_attachment=True,
                download_name=filename
            )
        else:
            # Use existing export for json/markdown
            files = sdlc_service.export_documents(srs, design, plan, format=export_format)
            return jsonify({
                'success': True,
                'files': files,
                'message': f'Documents exported in {export_format} format'
            })
        
    except Exception as e:
        logger.error(f"Error exporting SDLC documents: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

def generate_word_document(srs, design, ui_design, task_breakdowns, versioned_releases, plan, implementation_tools, test_plan, deployment_plan, maintenance_plan, project_title):
    """Generate Word document with all SDLC documents including task breakdown and versioned releases"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.paragraph_format.space_after = Pt(12)
    
    heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    heading_style.paragraph_format.space_after = Pt(6)
    
    # Title page
    doc.add_heading(f'{project_title} - Complete SDLC Documents', 0)
    doc.add_paragraph(f'Generated on {datetime.now().strftime("%B %d, %Y")}')
    doc.add_page_break()
    
    # Table of Contents placeholder
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('1. Software Requirements Specification (SRS)')
    doc.add_paragraph('2. System Design Document')
    doc.add_paragraph('3. UI Design Document')
    doc.add_paragraph('4. Implementation Plan')
    doc.add_paragraph('5. Implementation Tools and Resources')
    doc.add_paragraph('6. Task Breakdown')
    doc.add_paragraph('7. Versioned Product Releases')
    doc.add_paragraph('8. Test Plan')
    doc.add_paragraph('9. Deployment Plan')
    doc.add_paragraph('10. Maintenance Plan')
    doc.add_page_break()
    
    # SRS Section
    doc.add_heading('1. Software Requirements Specification (SRS)', level=1)
    doc.add_heading(f'{srs.project_title}', level=2)
    
    doc.add_heading('Overview', level=3)
    doc.add_paragraph(srs.overview or 'No overview provided.')
    
    doc.add_heading('Scope', level=3)
    doc.add_paragraph(srs.scope or 'No scope defined.')
    
    doc.add_heading('Functional Requirements', level=3)
    for req in srs.functional_requirements:
        p = doc.add_paragraph()
        p.add_run(f'{req["id"]}: ').bold = True
        p.add_run(req['description'])
    
    doc.add_heading('Non-Functional Requirements', level=3)
    for req in srs.non_functional_requirements:
        p = doc.add_paragraph()
        p.add_run(f'{req["id"]} ({req["category"]}): ').bold = True
        p.add_run(req['description'])
    
    doc.add_heading('User Stories', level=3)
    for story in srs.user_stories:
        p = doc.add_paragraph()
        p.add_run(f'{story["id"]}: ').bold = True
        p.add_run(story['story'])
    
    doc.add_heading('Acceptance Criteria', level=3)
    for criteria in srs.acceptance_criteria:
        doc.add_paragraph(criteria, style='List Bullet')
    
    doc.add_page_break()
    
    # Implementation Plan Section
    doc.add_heading('4. Implementation Plan', level=1)
    
    doc.add_heading('Timeline', level=3)
    timeline_text = f"Total Hours: {plan.timeline['total_hours']} | Total Days: {plan.timeline['total_days']} | Total Weeks: {plan.timeline['total_weeks']}"
    doc.add_paragraph(timeline_text)
    
    doc.add_heading('Phases', level=3)
    for phase in plan.phases:
        p = doc.add_paragraph()
        p.add_run(f'{phase["name"]}: ').bold = True
        p.add_run(f'{phase["hours"]} hours ({phase["days"]} days) - {phase["description"]}')
    
    doc.add_heading('Tasks', level=3)
    for task in plan.tasks:
        p = doc.add_paragraph()
        p.add_run(f'{task["title"]}: ').bold = True
        p.add_run(f'{task["estimated_hours"]} hours - {task["description"]}')
    
    doc.add_heading('Milestones', level=3)
    for milestone in plan.milestones:
        p = doc.add_paragraph()
        p.add_run(f'{milestone["name"]}: ').bold = True
        p.add_run(milestone['deliverable'])
    
    doc.add_heading('Resource Allocation', level=3)
    doc.add_paragraph(f"Developers needed: {plan.resource_allocation['developers_needed']}")
    doc.add_paragraph(f"Roles: {', '.join(plan.resource_allocation['roles'])}")
    
    # Implementation Tools Section
    doc.add_page_break()
    doc.add_heading('5. Implementation Tools and Resources', level=1)
    
    doc.add_heading('Development Tools', level=3)
    for tool in implementation_tools.get('development_tools', []):
        p = doc.add_paragraph()
        p.add_run(f'{tool["name"]}: ').bold = True
        p.add_run(f'{tool["description"]} - {tool["url"]}')
    
    doc.add_heading('Frameworks', level=3)
    for framework in implementation_tools.get('frameworks', []):
        p = doc.add_paragraph()
        p.add_run(f'{framework["name"]}: ').bold = True
        p.add_run(f'{framework["description"]} - {framework["url"]}')
    
    doc.add_heading('Databases', level=3)
    for db in implementation_tools.get('databases', []):
        p = doc.add_paragraph()
        p.add_run(f'{db["name"]}: ').bold = True
        p.add_run(f'{db["description"]} - {db["url"]}')
    
    doc.add_heading('Cloud Services', level=3)
    for service in implementation_tools.get('cloud_services', []):
        p = doc.add_paragraph()
        p.add_run(f'{service["name"]}: ').bold = True
        p.add_run(f'{service["description"]} - {service["url"]}')
    
    doc.add_heading('DevOps Tools', level=3)
    for tool in implementation_tools.get('devops_tools', []):
        p = doc.add_paragraph()
        p.add_run(f'{tool["name"]}: ').bold = True
        p.add_run(f'{tool["description"]} - {tool["url"]}')
    
    doc.add_heading('Testing Tools', level=3)
    for tool in implementation_tools.get('testing_tools', []):
        p = doc.add_paragraph()
        p.add_run(f'{tool["name"]}: ').bold = True
        p.add_run(f'{tool["description"]} - {tool["url"]}')
    
    doc.add_heading('Monitoring Tools', level=3)
    for tool in implementation_tools.get('monitoring_tools', []):
        p = doc.add_paragraph()
        p.add_run(f'{tool["name"]}: ').bold = True
        p.add_run(f'{tool["description"]} - {tool["url"]}')
    
    doc.add_heading('Security Tools', level=3)
    for tool in implementation_tools.get('security_tools', []):
        p = doc.add_paragraph()
        p.add_run(f'{tool["name"]}: ').bold = True
        p.add_run(f'{tool["description"]} - {tool["url"]}')
    
    doc.add_heading('Collaboration Tools', level=3)
    for tool in implementation_tools.get('collaboration_tools', []):
        p = doc.add_paragraph()
        p.add_run(f'{tool["name"]}: ').bold = True
        p.add_run(f'{tool["description"]} - {tool["url"]}')
    
    doc.add_heading('Learning Resources', level=3)
    for resource in implementation_tools.get('learning_resources', []):
        p = doc.add_paragraph()
        p.add_run(f'{resource["name"]}: ').bold = True
        p.add_run(f'{resource["description"]} - {resource["url"]}')
    
    # Task Breakdown Section
    doc.add_page_break()
    doc.add_heading('6. Task Breakdown', level=1)
    doc.add_paragraph(f'Total Requirements: {len(task_breakdowns)}')
    doc.add_paragraph(f'Total Tasks: {sum(len(bd["tasks"]) for bd in task_breakdowns)}')
    doc.add_paragraph(f'Total Estimated Hours: {sum(bd["estimated_hours"] for bd in task_breakdowns)}')
    
    for breakdown in task_breakdowns:
        doc.add_heading(f'{breakdown["requirement_id"]}: {breakdown["requirement_description"]}', level=2)
        doc.add_paragraph(f'Priority: {breakdown["priority"]} | Complexity: {breakdown["complexity"]} | Hours: {breakdown["estimated_hours"]}')
        
        doc.add_heading('Tasks:', level=3)
        for task in breakdown['tasks']:
            p = doc.add_paragraph()
            p.add_run(f'{task["title"]}').bold = True
            p.add_run(f' - {task["description"]} ({task["estimated_hours"]}h, {task["type"]}, {task["complexity"]})')
        
        if breakdown['dependencies']:
            doc.add_heading('Dependencies:', level=3)
            for dep in breakdown['dependencies']:
                doc.add_paragraph(f'• {dep}', style='List Bullet')
    
    doc.add_page_break()
    
    # Versioned Releases Section
    doc.add_heading('7. Versioned Product Releases', level=1)
    doc.add_paragraph(f'Total Versions: {len(versioned_releases)}')
    doc.add_paragraph(f'Total Features: {sum(len(vr["features"]) for vr in versioned_releases)}')
    doc.add_paragraph(f'Total Tasks: {sum(len(vr["tasks"]) for vr in versioned_releases)}')
    
    for release in versioned_releases:
        doc.add_heading(f'{release["version"]} - {release["name"]}', level=2)
        doc.add_paragraph(f'Estimated Hours: {release["estimated_hours"]}')
        doc.add_paragraph(release['description'])
        
        doc.add_heading('Features:', level=3)
        for feature in release['features']:
            p = doc.add_paragraph()
            p.add_run(f'{feature["name"]}').bold = True
            p.add_run(f' - {feature["description"]} ({feature["type"]})')
        
        doc.add_heading('Tasks:', level=3)
        for task in release['tasks']:
            p = doc.add_paragraph()
            p.add_run(f'{task["title"]}').bold = True
            p.add_run(f' - {task["estimated_hours"]}h ({task["type"]})')
        
        if release['dependencies']:
            doc.add_heading('Dependencies:', level=3)
            for dep in release['dependencies']:
                doc.add_paragraph(f'• {dep}', style='List Bullet')
        
        doc.add_heading('Release Criteria:', level=3)
        for criteria in release['release_criteria']:
            doc.add_paragraph(f'• {criteria}', style='List Bullet')
        
        doc.add_heading('Testing Requirements:', level=3)
        for req in release['testing_requirements']:
            doc.add_paragraph(f'• {req}', style='List Bullet')
        
        doc.add_heading('Deployment Notes:', level=3)
        doc.add_paragraph(release['deployment_notes'])
    
    doc.add_page_break()
    
    # Design Section
    doc.add_heading('2. System Design Document', level=1)
    
    doc.add_heading('Architecture Type', level=3)
    doc.add_paragraph(design.architecture_type)
    
    doc.add_heading('Components', level=3)
    for comp in design.components:
        p = doc.add_paragraph()
        p.add_run(f'{comp["name"]}: ').bold = True
        p.add_run(comp['description'])
    
    doc.add_heading('Data Models', level=3)
    for model in design.data_models:
        p = doc.add_paragraph()
        p.add_run(f'{model["name"]}: ').bold = True
        p.add_run(f'Fields: {model["fields"]}, Relationships: {model["relationships"]}')
    
    doc.add_heading('API Endpoints', level=3)
    for endpoint in design.api_endpoints:
        p = doc.add_paragraph()
        p.add_run(f'{endpoint["method"]} {endpoint["path"]}: ').bold = True
        p.add_run(endpoint['description'])
    
    doc.add_heading('Technology Stack', level=3)
    for category, techs in design.technology_stack.items():
        p = doc.add_paragraph()
        p.add_run(f'{category}: ').bold = True
        p.add_run(', '.join(techs))
    
    doc.add_heading('Security Considerations', level=3)
    for security in design.security_considerations:
        doc.add_paragraph(security, style='List Bullet')
    
    doc.add_page_break()
    
    # UI Design Section
    doc.add_heading('3. UI Design Document', level=1)
    
    doc.add_heading('Design System', level=3)
    for key, value in ui_design.design_system.items():
        p = doc.add_paragraph()
        p.add_run(f'{key}: ').bold = True
        p.add_run(str(value))
    
    doc.add_heading('Page Layouts', level=3)
    for layout in ui_design.page_layouts:
        p = doc.add_paragraph()
        p.add_run(f'{layout["name"]}: ').bold = True
        p.add_run(f'{layout["type"]} - {layout["description"]}')
    
    doc.add_heading('UI Components', level=3)
    for component in ui_design.components:
        p = doc.add_paragraph()
        p.add_run(f'{component["name"]}: ').bold = True
        p.add_run(f'{component["type"]} - {component["description"]}')
    
    doc.add_heading('Wireframes', level=3)
    for wireframe in ui_design.wireframes:
        p = doc.add_paragraph()
        p.add_run(f'{wireframe["page"]}: ').bold = True
        p.add_run(f'{wireframe["layout_type"]} - {wireframe["description"]}')
    
    doc.add_heading('Responsive Breakpoints', level=3)
    for breakpoint in ui_design.responsive_breakpoints:
        doc.add_paragraph(breakpoint, style='List Bullet')
    
    doc.add_heading('Accessibility Features', level=3)
    for feature in ui_design.accessibility_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('CSS Framework', level=3)
    doc.add_paragraph(ui_design.css_framework)
    
    doc.add_heading('JavaScript Libraries', level=3)
    for library in ui_design.javascript_libraries:
        doc.add_paragraph(library, style='List Bullet')
    
    doc.add_heading('Interactive Elements', level=3)
    for element in ui_design.interactive_elements:
        p = doc.add_paragraph()
        p.add_run(f'{element["name"]}: ').bold = True
        p.add_run(f'{element["type"]} - {element["description"]}')
    
    doc.add_page_break()
    
    # Test Plan Section
    doc.add_heading('8. Test Plan', level=1)
    doc.add_heading('Test Strategy', level=3)
    doc.add_paragraph(test_plan.get('test_strategy', 'No test strategy defined.'))
    
    doc.add_heading('Test Phases', level=3)
    for phase in test_plan.get('test_phases', []):
        p = doc.add_paragraph()
        p.add_run(f'{phase["phase"]}: ').bold = True
        p.add_run(phase['description'])
    
    doc.add_heading('Test Environment', level=3)
    doc.add_paragraph(test_plan.get('test_environment', 'No test environment defined.'))
    
    doc.add_heading('Automation Plan', level=3)
    doc.add_paragraph(test_plan.get('automation_plan', 'No automation plan defined.'))
    
    # Deployment Plan Section
    doc.add_page_break()
    doc.add_heading('9. Deployment Plan', level=1)
    doc.add_heading('Deployment Strategy', level=3)
    doc.add_paragraph(deployment_plan.get('deployment_strategy', 'No deployment strategy defined.'))
    
    doc.add_heading('Environments', level=3)
    for env in deployment_plan.get('environments', []):
        doc.add_paragraph(f'• {env}', style='List Bullet')
    
    doc.add_heading('Deployment Steps', level=3)
    for step in deployment_plan.get('deployment_steps', []):
        doc.add_paragraph(f'• {step}', style='List Bullet')
    
    doc.add_heading('Rollback Plan', level=3)
    doc.add_paragraph(deployment_plan.get('rollback_plan', 'No rollback plan defined.'))
    
    doc.add_heading('Monitoring', level=3)
    doc.add_paragraph(deployment_plan.get('monitoring', 'No monitoring plan defined.'))
    
    # Maintenance Plan Section
    doc.add_page_break()
    doc.add_heading('10. Maintenance Plan', level=1)
    doc.add_heading('Maintenance Schedule', level=3)
    doc.add_paragraph(maintenance_plan.get('maintenance_schedule', 'No maintenance schedule defined.'))
    
    doc.add_heading('Monitoring', level=3)
    doc.add_paragraph(maintenance_plan.get('monitoring', 'No monitoring plan defined.'))
    
    doc.add_heading('Backup Strategy', level=3)
    doc.add_paragraph(maintenance_plan.get('backup_strategy', 'No backup strategy defined.'))
    
    doc.add_heading('Security Updates', level=3)
    doc.add_paragraph(maintenance_plan.get('security_updates', 'No security update plan defined.'))
    
    doc.add_heading('Performance Optimization', level=3)
    doc.add_paragraph(maintenance_plan.get('performance_optimization', 'No performance optimization plan defined.'))
    
    doc.add_heading('Support Plan', level=3)
    doc.add_paragraph(maintenance_plan.get('support_plan', 'No support plan defined.'))
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer

def generate_pdf_document(srs, design, ui_design, task_breakdowns, versioned_releases, plan, implementation_tools, test_plan, deployment_plan, maintenance_plan, project_title):
    """Generate PDF document with all SDLC documents including task breakdown and versioned releases"""
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        alignment=1  # Center
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=6
    )
    
    story = []
    
    # Title page
    story.append(Paragraph(f'{project_title} - Complete SDLC Documents', title_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f'Generated on {datetime.now().strftime("%B %d, %Y")}', styles['Normal']))
    story.append(PageBreak())
    
    # Table of Contents
    story.append(Paragraph('Table of Contents', heading_style))
    story.append(Paragraph('1. Software Requirements Specification (SRS)', styles['Normal']))
    story.append(Paragraph('2. System Design Document', styles['Normal']))
    story.append(Paragraph('3. UI Design Document', styles['Normal']))
    story.append(Paragraph('4. Implementation Plan', styles['Normal']))
    story.append(Paragraph('5. Implementation Tools and Resources', styles['Normal']))
    story.append(Paragraph('6. Task Breakdown', styles['Normal']))
    story.append(Paragraph('7. Versioned Product Releases', styles['Normal']))
    story.append(Paragraph('8. Test Plan', styles['Normal']))
    story.append(Paragraph('9. Deployment Plan', styles['Normal']))
    story.append(Paragraph('10. Maintenance Plan', styles['Normal']))
    story.append(PageBreak())
    
    # SRS Section
    story.append(Paragraph('1. Software Requirements Specification (SRS)', heading_style))
    story.append(Paragraph(f'{srs.project_title}', styles['Heading2']))
    
    story.append(Paragraph('Overview', styles['Heading3']))
    story.append(Paragraph(srs.overview or 'No overview provided.', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Scope', styles['Heading3']))
    story.append(Paragraph(srs.scope or 'No scope defined.', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Functional Requirements', styles['Heading3']))
    for req in srs.functional_requirements:
        story.append(Paragraph(f'<b>{req["id"]}:</b> {req["description"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Non-Functional Requirements', styles['Heading3']))
    for req in srs.non_functional_requirements:
        story.append(Paragraph(f'<b>{req["id"]} ({req["category"]}):</b> {req["description"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('User Stories', styles['Heading3']))
    for story_item in srs.user_stories:
        story.append(Paragraph(f'<b>{story_item["id"]}:</b> {story_item["story"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Acceptance Criteria', styles['Heading3']))
    for criteria in srs.acceptance_criteria:
        story.append(Paragraph(f'• {criteria}', styles['Normal']))
    story.append(PageBreak())
    
    # Implementation Plan Section
    story.append(Paragraph('4. Implementation Plan', heading_style))
    
    story.append(Paragraph('Timeline', styles['Heading3']))
    timeline_text = f"Total Hours: {plan.timeline['total_hours']} | Total Days: {plan.timeline['total_days']} | Total Weeks: {plan.timeline['total_weeks']}"
    story.append(Paragraph(timeline_text, styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Phases', styles['Heading3']))
    for phase in plan.phases:
        story.append(Paragraph(f'<b>{phase["name"]}:</b> {phase["hours"]} hours ({phase["days"]} days) - {phase["description"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Tasks', styles['Heading3']))
    for task in plan.tasks:
        story.append(Paragraph(f'<b>{task["title"]}:</b> {task["estimated_hours"]} hours - {task["description"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Milestones', styles['Heading3']))
    for milestone in plan.milestones:
        story.append(Paragraph(f'<b>{milestone["name"]}:</b> {milestone["deliverable"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Resource Allocation', styles['Heading3']))
    story.append(Paragraph(f"Developers needed: {plan.resource_allocation['developers_needed']}", styles['Normal']))
    story.append(Paragraph(f"Roles: {', '.join(plan.resource_allocation['roles'])}", styles['Normal']))
    story.append(PageBreak())
    
    # Implementation Tools Section
    story.append(Paragraph('5. Implementation Tools and Resources', heading_style))
    
    story.append(Paragraph('Development Tools', styles['Heading3']))
    for tool in implementation_tools.get('development_tools', []):
        story.append(Paragraph(f'<b>{tool["name"]}:</b> {tool["description"]} - {tool["url"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Frameworks', styles['Heading3']))
    for framework in implementation_tools.get('frameworks', []):
        story.append(Paragraph(f'<b>{framework["name"]}:</b> {framework["description"]} - {framework["url"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Databases', styles['Heading3']))
    for db in implementation_tools.get('databases', []):
        story.append(Paragraph(f'<b>{db["name"]}:</b> {db["description"]} - {db["url"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Cloud Services', styles['Heading3']))
    for service in implementation_tools.get('cloud_services', []):
        story.append(Paragraph(f'<b>{service["name"]}:</b> {service["description"]} - {service["url"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('DevOps Tools', styles['Heading3']))
    for tool in implementation_tools.get('devops_tools', []):
        story.append(Paragraph(f'<b>{tool["name"]}:</b> {tool["description"]} - {tool["url"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Testing Tools', styles['Heading3']))
    for tool in implementation_tools.get('testing_tools', []):
        story.append(Paragraph(f'<b>{tool["name"]}:</b> {tool["description"]} - {tool["url"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Monitoring Tools', styles['Heading3']))
    for tool in implementation_tools.get('monitoring_tools', []):
        story.append(Paragraph(f'<b>{tool["name"]}:</b> {tool["description"]} - {tool["url"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Security Tools', styles['Heading3']))
    for tool in implementation_tools.get('security_tools', []):
        story.append(Paragraph(f'<b>{tool["name"]}:</b> {tool["description"]} - {tool["url"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Collaboration Tools', styles['Heading3']))
    for tool in implementation_tools.get('collaboration_tools', []):
        story.append(Paragraph(f'<b>{tool["name"]}:</b> {tool["description"]} - {tool["url"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Learning Resources', styles['Heading3']))
    for resource in implementation_tools.get('learning_resources', []):
        story.append(Paragraph(f'<b>{resource["name"]}:</b> {resource["description"]} - {resource["url"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Task Breakdown Section
    story.append(Paragraph('6. Task Breakdown', heading_style))
    story.append(Paragraph(f'Total Requirements: {len(task_breakdowns)}', styles['Normal']))
    story.append(Paragraph(f'Total Tasks: {sum(len(bd["tasks"]) for bd in task_breakdowns)}', styles['Normal']))
    story.append(Paragraph(f'Total Estimated Hours: {sum(bd["estimated_hours"] for bd in task_breakdowns)}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    for breakdown in task_breakdowns:
        story.append(Paragraph(f'<b>{breakdown["requirement_id"]}: {breakdown["requirement_description"]}</b>', styles['Heading3']))
        story.append(Paragraph(f'Priority: {breakdown["priority"]} | Complexity: {breakdown["complexity"]} | Hours: {breakdown["estimated_hours"]}', styles['Normal']))
        story.append(Spacer(1, 6))
        
        story.append(Paragraph('Tasks:', styles['Heading4']))
        for task in breakdown['tasks']:
            story.append(Paragraph(f'• <b>{task["title"]}</b> - {task["description"]} ({task["estimated_hours"]}h, {task["type"]}, {task["complexity"]})', styles['Normal']))
        
        if breakdown['dependencies']:
            story.append(Paragraph('Dependencies:', styles['Heading4']))
            for dep in breakdown['dependencies']:
                story.append(Paragraph(f'• {dep}', styles['Normal']))
        
        story.append(Spacer(1, 12))
    
    story.append(PageBreak())
    
    # Versioned Releases Section
    story.append(Paragraph('7. Versioned Product Releases', heading_style))
    story.append(Paragraph(f'Total Versions: {len(versioned_releases)}', styles['Normal']))
    story.append(Paragraph(f'Total Features: {sum(len(vr["features"]) for vr in versioned_releases)}', styles['Normal']))
    story.append(Paragraph(f'Total Tasks: {sum(len(vr["tasks"]) for vr in versioned_releases)}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    for release in versioned_releases:
        story.append(Paragraph(f'<b>{release["version"]} - {release["name"]}</b>', styles['Heading3']))
        story.append(Paragraph(f'Estimated Hours: {release["estimated_hours"]}', styles['Normal']))
        story.append(Paragraph(release['description'], styles['Normal']))
        story.append(Spacer(1, 6))
        
        story.append(Paragraph('Features:', styles['Heading4']))
        for feature in release['features']:
            story.append(Paragraph(f'• <b>{feature["name"]}</b> - {feature["description"]} ({feature["type"]})', styles['Normal']))
        
        story.append(Paragraph('Tasks:', styles['Heading4']))
        for task in release['tasks']:
            story.append(Paragraph(f'• <b>{task["title"]}</b> - {task["estimated_hours"]}h ({task["type"]})', styles['Normal']))
        
        if release['dependencies']:
            story.append(Paragraph('Dependencies:', styles['Heading4']))
            for dep in release['dependencies']:
                story.append(Paragraph(f'• {dep}', styles['Normal']))
        
        story.append(Paragraph('Release Criteria:', styles['Heading4']))
        for criteria in release['release_criteria']:
            story.append(Paragraph(f'• {criteria}', styles['Normal']))
        
        story.append(Paragraph('Testing Requirements:', styles['Heading4']))
        for req in release['testing_requirements']:
            story.append(Paragraph(f'• {req}', styles['Normal']))
        
        story.append(Paragraph('Deployment Notes:', styles['Heading4']))
        story.append(Paragraph(release['deployment_notes'], styles['Normal']))
        story.append(Spacer(1, 12))
    
    story.append(PageBreak())
    
    # Design Section
    story.append(Paragraph('2. System Design Document', heading_style))
    
    story.append(Paragraph('Architecture Type', styles['Heading3']))
    story.append(Paragraph(design.architecture_type, styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Components', styles['Heading3']))
    for comp in design.components:
        story.append(Paragraph(f'<b>{comp["name"]}:</b> {comp["description"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Data Models', styles['Heading3']))
    for model in design.data_models:
        story.append(Paragraph(f'<b>{model["name"]}:</b> Fields: {model["fields"]}, Relationships: {model["relationships"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('API Endpoints', styles['Heading3']))
    for endpoint in design.api_endpoints:
        story.append(Paragraph(f'<b>{endpoint["method"]} {endpoint["path"]}:</b> {endpoint["description"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Technology Stack', styles['Heading3']))
    for category, techs in design.technology_stack.items():
        story.append(Paragraph(f'<b>{category}:</b> {", ".join(techs)}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Security Considerations', styles['Heading3']))
    for security in design.security_considerations:
        story.append(Paragraph(f'• {security}', styles['Normal']))
    story.append(PageBreak())
    
    # UI Design Section
    story.append(Paragraph('3. UI Design Document', heading_style))
    
    story.append(Paragraph('Design System', styles['Heading3']))
    for key, value in ui_design.design_system.items():
        story.append(Paragraph(f'<b>{key}:</b> {str(value)}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Page Layouts', styles['Heading3']))
    for layout in ui_design.page_layouts:
        story.append(Paragraph(f'<b>{layout["name"]}:</b> {layout["type"]} - {layout["description"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('UI Components', styles['Heading3']))
    for component in ui_design.components:
        story.append(Paragraph(f'<b>{component["name"]}:</b> {component["type"]} - {component["description"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Wireframes', styles['Heading3']))
    for wireframe in ui_design.wireframes:
        story.append(Paragraph(f'<b>{wireframe["page"]}:</b> {wireframe["layout_type"]} - {wireframe["description"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Responsive Breakpoints', styles['Heading3']))
    for breakpoint in ui_design.responsive_breakpoints:
        story.append(Paragraph(f'• {breakpoint}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Accessibility Features', styles['Heading3']))
    for feature in ui_design.accessibility_features:
        story.append(Paragraph(f'• {feature}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('CSS Framework', styles['Heading3']))
    story.append(Paragraph(ui_design.css_framework, styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('JavaScript Libraries', styles['Heading3']))
    for library in ui_design.javascript_libraries:
        story.append(Paragraph(f'• {library}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Interactive Elements', styles['Heading3']))
    for element in ui_design.interactive_elements:
        story.append(Paragraph(f'• {element["name"]}: {element["type"]} - {element["description"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Test Plan Section
    story.append(Paragraph('8. Test Plan', heading_style))
    story.append(Paragraph('Test Strategy', styles['Heading3']))
    story.append(Paragraph(test_plan.get('test_strategy', 'No test strategy defined.'), styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Test Phases', styles['Heading3']))
    for phase in test_plan.get('test_phases', []):
        story.append(Paragraph(f'<b>{phase["phase"]}:</b> {phase["description"]}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Test Environment', styles['Heading3']))
    story.append(Paragraph(test_plan.get('test_environment', 'No test environment defined.'), styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Automation Plan', styles['Heading3']))
    story.append(Paragraph(test_plan.get('automation_plan', 'No automation plan defined.'), styles['Normal']))
    story.append(PageBreak())
    
    # Deployment Plan Section
    story.append(Paragraph('9. Deployment Plan', heading_style))
    story.append(Paragraph('Deployment Strategy', styles['Heading3']))
    story.append(Paragraph(deployment_plan.get('deployment_strategy', 'No deployment strategy defined.'), styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Environments', styles['Heading3']))
    for env in deployment_plan.get('environments', []):
        story.append(Paragraph(f'• {env}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Deployment Steps', styles['Heading3']))
    for step in deployment_plan.get('deployment_steps', []):
        story.append(Paragraph(f'• {step}', styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Rollback Plan', styles['Heading3']))
    story.append(Paragraph(deployment_plan.get('rollback_plan', 'No rollback plan defined.'), styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Monitoring', styles['Heading3']))
    story.append(Paragraph(deployment_plan.get('monitoring', 'No monitoring plan defined.'), styles['Normal']))
    story.append(PageBreak())
    
    # Maintenance Plan Section
    story.append(Paragraph('10. Maintenance Plan', heading_style))
    story.append(Paragraph('Maintenance Schedule', styles['Heading3']))
    story.append(Paragraph(maintenance_plan.get('maintenance_schedule', 'No maintenance schedule defined.'), styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Monitoring', styles['Heading3']))
    story.append(Paragraph(maintenance_plan.get('monitoring', 'No monitoring plan defined.'), styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Backup Strategy', styles['Heading3']))
    story.append(Paragraph(maintenance_plan.get('backup_strategy', 'No backup strategy defined.'), styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Security Updates', styles['Heading3']))
    story.append(Paragraph(maintenance_plan.get('security_updates', 'No security update plan defined.'), styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Performance Optimization', styles['Heading3']))
    story.append(Paragraph(maintenance_plan.get('performance_optimization', 'No performance optimization plan defined.'), styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Support Plan', styles['Heading3']))
    story.append(Paragraph(maintenance_plan.get('support_plan', 'No support plan defined.'), styles['Normal']))
    
    # Build PDF
    doc.build(story)
    return buffer

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    host = os.environ.get('HOST', '127.0.0.1')
    
    logger.info(f"Starting Project Management API on {host}:{port}")
    logger.info(f"Dashboard will be available at: http://{host}:{port}")
    
    app.run(host=host, port=port, debug=True) 